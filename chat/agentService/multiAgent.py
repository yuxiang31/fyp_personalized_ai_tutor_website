from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_tavily import TavilySearch
from langgraph_supervisor import create_supervisor
from langgraph.prebuilt import create_react_agent
from langchain_core.messages import AIMessage, HumanMessage, ToolMessage
from langchain_core.tools import tool

from django.conf import settings

from langgraph.checkpoint.postgres.aio import AsyncPostgresSaver
from psycopg_pool import AsyncConnectionPool

from dotenv import load_dotenv

from datetime import datetime
from rich.console import Console
from rich.panel import Panel

import re
from docx import Document  
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import os
from typing import List, Optional

import logging
from asgiref.sync import sync_to_async

logger = logging.getLogger(__name__)
# Configure logger to show more detailed information
logging.basicConfig(level=logging.INFO)

load_dotenv()
console = Console()
_RICH_AVAILABLE = True

############################################################### LLM DEFINITION SECTION ###############################################################
'''
Define and configure LLM used
'''

# LLM Definition
llm = ChatGoogleGenerativeAI(
    model="gemini-2.0-flash",
    google_api_key=os.getenv("GEMINI_API_KEY"),
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
)

############################################################### LLM DEFINITION SECTION ###############################################################

# under async method
# 1. LLM definition
# 2. Database URL definition
# 3. Tools definition
# 4. Create a database connection pool
# 5. Agent Definition (single agent -> learning materials agent, planner agent)
#     5.1 Implement llmt
#     5.2 Implement tools
#     5.3 Implement prompt
#     5.4 Implement name
# 6. Implement multi-agent collaboration
    # 6.1 Implement supervisor agent
    # 6.2 Implement memory using data connection pool
    # 6.3 Supervisor agent will be mainly interact with user message
    # 6.4 Collaboration between agent will not show in chat page
# 7. Able to integrate to consumers.py 

############################################################### PYTHON FUNCTION DEFINITION ###############################################################
'''
Additional Function Definition for Tools
'''
def set_table_borders(table):
    # Access the table's XML element
    tbl = table._tbl

    # Create or get <w:tblPr>
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')

    # Create border elements
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')       # Type of border
        border.set(qn('w:sz'), '4')             # Width (4 eighths of a point)
        border.set(qn('w:space'), '0')          # No spacing
        border.set(qn('w:color'), '000000')     # Black color
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)

# use to parse text generated from learning materials agent
def parse_text(text):
    result = {
        'topic': '',
        'overview': '',
        'sub_topics': [],
        'references': [],
    }

    # Extract topic
    topic_match = re.search(r'<topic>\s*Topic:\s*(.*?)\s*</topic>', text, re.DOTALL)
    if topic_match:
        result['topic'] = topic_match.group(1).strip()

    # Extract overview
    overview_match = re.search(r'<overview>(.*?)</overview>', text, re.DOTALL)
    if overview_match:
        result['overview'] = overview_match.group(1).strip()

    # Extract sub-topics and paragraphs
    sub_topic_pattern = re.compile(r'<sub-topic>\s*(.*?)\s*</sub-topic>\s*((?:<paragraph>.*?</paragraph>\s*)+)', re.DOTALL)
    paragraph_pattern = re.compile(r'<paragraph>\s*(.*?)\s*</paragraph>', re.DOTALL)

    # Extract URLs
    urls_result = []
    urls_match = re.search(r'<references>(.*?)</references>', text, re.DOTALL)
    if urls_match:
        urls_block = urls_match.group(1).strip()
        if urls_block:
            urls_result = [u.strip() for u in urls_block.split("\n") if u.strip()]

    for url in urls_result:
        result['references'].append(url)

    for match in sub_topic_pattern.finditer(text):
        sub_topic_title = match.group(1).strip()
        paragraphs_block = match.group(2)

        paragraphs = [p.strip() for p in paragraph_pattern.findall(paragraphs_block)]
        result['sub_topics'].append({
            'title': sub_topic_title,
            'paragraphs': paragraphs
        })

    return result
############################################################### PYTHON FUNCTION DEFINITION ###############################################################


############################################################### TOOLS DEFINITION ###############################################################
'''
Define tools for AI Agent
1. Draft Learning Plan Tool (Draft Learning Plan Tool for Planner Agent)
2. Tavily Web Search Tool (Web Search Tool for Learning Materials Agent)
3. Draft Learning Materials Tool (Draft Learning Materials Tool for Learning Materials Agent)
'''

@tool
def _web_search(query: str):
    """Tool that able to perform web search and get latest and relevant results"""
    web_search = TavilySearch(max_results=3)
    web_search_results = web_search.invoke(query)

    return web_search_results["results"]

@tool("write_study_plan")
def _write_study_plan(plan: str, filename: Optional[str] = None) -> str:
    """Persist a study plan: 
    1) save the raw plan into DB, 
    2) write a .docx under static/generated/study_plans, and 
    3) return a JSON string with download_url and IDs. 
    Always pass the full plan text as the 'plan' argument.
    """

    # Local import to avoid circular imports at module import time
    try:
        from chat.models import StudyMaterials, StudyMaterialFiles
    except Exception as e:
        return f"ERROR: unable to import models - {e}"

    if not plan or not isinstance(plan, str) or plan.strip() == "":
        return "ERROR: 'plan' must be a non-empty string."

    def extract_tag_content(text: str, tag: str):
        pattern = fr"<{tag}>\s*(.*?)\s*</{tag}>"
        return re.findall(pattern, text, re.DOTALL)

    def parse_and_write_to_word(paragraph_text: str, output_path: str):
        doc = Document()

        # Extract subject name
        subject_matches = extract_tag_content(paragraph_text, "subject name")
        subject_name = subject_matches[0].strip() if subject_matches else "Study Plan"
        doc.add_heading(f"{subject_name}", level=1)

        # Process each <topic X> section
        topic_sections = re.findall(r"<topic \d+>(.*?)</topic \d+>", paragraph_text, re.DOTALL)
        
        for idx, topic_content in enumerate(topic_sections, start=1):
            # Extract topic name and overview
            topic_name_matches = extract_tag_content(topic_content, "Topic Name")
            topic_overview_matches = extract_tag_content(topic_content, "Topic Overview")
            topic_name = topic_name_matches[0].strip() if topic_name_matches else f"Topic {idx}"
            topic_overview = topic_overview_matches[0].strip() if topic_overview_matches else ""

            # Add topic heading and overview
            doc.add_heading(f"Topic: {topic_name}", level=1)
            overview_paragraph = doc.add_paragraph(f"Overview: {topic_overview}", style='Normal')
            overview_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Add table with headers
            table = doc.add_table(rows=1, cols=5)
            set_table_borders(table)  # Add borders to the table

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No'
            hdr_cells[1].text = 'Sub-Topic'
            hdr_cells[2].text = 'Study Objectives'
            hdr_cells[3].text = 'Estimated Time'
            hdr_cells[4].text = 'Checkbox'

            # Extract sub-topic details
            sub_topics = extract_tag_content(topic_content, "Sub-Topic Name")
            objectives = extract_tag_content(topic_content, "Study-Objectives")
            times = extract_tag_content(topic_content, "Estimated time")

            for i in range(len(sub_topics)):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = sub_topics[i].strip()
                row_cells[2].text = objectives[i].strip() if i < len(objectives) else ""
                row_cells[3].text = times[i].strip() if i < len(times) else ""
                row_cells[4].text = ""  # Checkbox column

            doc.add_paragraph()  # Add space before next topic

        # Save the document
        doc.save(output_path)

    # 1) Store plan as text in DB
    try:
        material = StudyMaterials.objects.create(study_plan=plan)
    except Exception as e:
        return f"ERROR: failed to store study plan in DB - {e}"

    # 2) Prepare filesystem paths under static/generated/study_plans
    try:
        base_dir = settings.BASE_DIR
        static_url = settings.STATIC_URL
        out_dir = os.path.join(base_dir, "static", "generated", "study_plans")
        os.makedirs(out_dir, exist_ok=True)

        # Prefer subject name if present; otherwise use DB-generated ID
        subject_matches = extract_tag_content(plan, "subject name")
        base_name = subject_matches[0].strip() if subject_matches else f"study-plan-{material.id}"
        # Sanitize filename (remove problematic chars)
        safe_base = re.sub(r"[^A-Za-z0-9\-_ ]+", "", base_name).strip().replace(" ", "-") or f"study-plan-{material.id}"
        file_name = safe_base + ".docx"
        output_path = os.path.join(out_dir, file_name)

        # Write the .docx
        parse_and_write_to_word(plan, output_path)

        # Store file reference in DB (relative path for portability)
        rel_path = os.path.join("static", "generated", "study_plans", file_name).replace("\\", "/")
        StudyMaterialFiles.objects.create(study_material=material, file_path=rel_path)

        # Build download URL (ensure leading slash)
        static_prefix = f"/{static_url.strip('/')}" if not static_url.startswith("/") else static_url
        download_url = f"{static_prefix.rstrip('/')}/generated/study_plans/{file_name}"

        return (
            "{"
            f"\"study_material_id\": \"{material.id}\", "
            f"\"file_path\": \"{rel_path}\", "
            f"\"download_url\": \"{download_url}\""
            "}"
        )
    except Exception as e:
        return f"ERROR: failed to write study plan file - {e}"

@tool("draft_notes")
def _draft_notes(contents: str, filename: Optional[str] = None) -> str:
    """
       USE TO RECORD CONTENTS.
       Tool that able to write the content generated and generate a word document notes. It contain three sections which are
       topic, overview section, and sub-topic section.
       topic section - Name of the topic
       overview section - Overview of the topic given
       sub-topics section - Each sub-topic will have explanations provided
    """

    content_dict = parse_text(contents)

    # Assign variables with safe defaults
    topic = (content_dict.get('topic') or 'Study Notes').strip()
    overview = (content_dict.get('overview') or '').strip()
    sub_topics = content_dict.get('sub_topics') or []  # list of {title, paragraphs}
    references = content_dict.get('references') or []

    # Build the document
    document = Document()
    document.add_heading(f'Topic: {topic}', level=0)
    document.add_heading('Overview', level=1)
    if overview:
        overview_paragraph = document.add_paragraph(overview)
        overview_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for sub_topic in sub_topics:
        title = sub_topic.get('title', '').strip()
        paragraphs = sub_topic.get('paragraphs', []) or []
        if title:
            document.add_heading(f"{title}", level=1)
        if paragraphs:
            # First paragraph
            p = document.add_paragraph(paragraphs[0])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remaining paragraphs, each as its own paragraph
            for para in paragraphs[1:]:
                p_next = document.add_paragraph(para)
                p_next.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if references:
        document.add_heading("References", level=1)
        for reference in references:
            if reference:
                document.add_paragraph(reference, style='ListBullet')

    # Prepare output path under static/generated/study_materials
    try:
        base_dir = settings.BASE_DIR
        static_url = settings.STATIC_URL
        out_dir = os.path.join(base_dir, 'static', 'generated', 'study_materials')
        os.makedirs(out_dir, exist_ok=True)

        # File name: prefer provided filename, else topic
        base_name = (filename or topic or 'study-notes').strip()
        safe_base = re.sub(r"[^A-Za-z0-9\-_ ]+", "", base_name).strip().replace(" ", "-") or 'study-notes'
        file_name = f"{safe_base}.docx"
        output_path = os.path.join(out_dir, file_name)

        # Save document
        document.save(output_path)

        # Relative path and download URL
        rel_path = os.path.join('static', 'generated', 'study_materials', file_name).replace('\\', '/')
        static_prefix = f"/{static_url.strip('/')}" if not static_url.startswith('/') else static_url
        download_url = f"{static_prefix.rstrip('/')}/generated/study_materials/{file_name}"

        # Best-effort: attach this file to the current conversation's StudyMaterials
        study_material_id_str: Optional[str] = None
        try:
            from chat.models import StudyMaterialFiles, ConversationRecord, Checkpoints
            conversation = None
            # Prefer resolving via the latest checkpoint's thread_id
            try:
                latest_cp = Checkpoints.objects.order_by('-id').first()
            except Exception:
                latest_cp = None
            if latest_cp and latest_cp.thread_id:
                try:
                    conversation = ConversationRecord.objects.filter(thread_id=latest_cp.thread_id).first()
                except Exception:
                    conversation = None
            # Fallback: latest updated conversation
            if conversation is None:
                try:
                    conversation = ConversationRecord.objects.order_by('-updated_at').first()
                except Exception:
                    conversation = None

            if conversation and conversation.study_material_id:
                StudyMaterialFiles.objects.create(
                    study_material_id=conversation.study_material_id,
                    file_path=rel_path,
                )
                study_material_id_str = str(conversation.study_material_id)
        except Exception:
            # Do not fail the tool if DB linkage is not possible
            pass

        # Return minimal JSON-like string for the agent tool result
        return (
            "{"
            f"\"topic\": \"{topic}\", "
            f"\"file_path\": \"{rel_path}\", "
            f"\"download_url\": \"{download_url}\"" +
            (f", \"study_material_id\": \"{study_material_id_str}\"" if study_material_id_str else "") +
            "}"
        )
    except Exception as e:
        return f"ERROR: failed to write notes file - {e}"
############################################################### TOOLS DEFINITION ###############################################################

############################################################### MULTI AGENT COLLABORATION CONSOLE MONITORING ###############################################################
'''
Debug printing for agent collaboration using Rich
This lives in the multi-agent collaboration section so you can observe all cross-agent traffic.
'''

_AGENT_COLOR_MAP = {
    "user": "bright_green",
    "supervisor": "bright_cyan",
    "maths_agent": "bright_yellow",
    "researcher_agent": "bright_magenta",
    "tool": "bright_blue",
    "tool_call": "bright_blue",
    "tool_result": "bright_blue",
    "system": "bright_black",
}

def _extract_text_content(content) -> str:
    """Normalize LC message content which may be str or list[dict/str] into plain text."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for c in content:
            if isinstance(c, dict):
                # Common LC content blocks have a `text` key
                txt = c.get("text") or c.get("content") or ""
                if isinstance(txt, list):
                    parts.extend([str(t) for t in txt])
                elif txt:
                    parts.append(str(txt))
            elif isinstance(c, str):
                parts.append(c)
        return "\n".join(p for p in parts if p)
    # Fallback
    try:
        return str(content)
    except Exception:
        return ""

def _print_collab_message(node_name: str, message) -> None:
    """Pretty-print a single collaboration message to the console using Rich.

    node_name: the graph node emitting the message (e.g., 'learning_materials_agent', 'planner_agent', 'supervisor').
    message: a LangChain message (HumanMessage, AIMessage, etc.).
    """
    # Prefer the LC message.name (often the agent name) when present
    msg_name = getattr(message, "name", None)
    role = (
        "user" if isinstance(message, HumanMessage) else
        (msg_name or node_name or "assistant")
    )
    color = _AGENT_COLOR_MAP.get(role, "white")
    title = f"{role.upper()}" if role in ("user", "supervisor") else f"{role.upper()} [{node_name}]"

    body = _extract_text_content(getattr(message, "content", ""))
    timestamp = datetime.now().strftime("%H:%M:%S")

    if _RICH_AVAILABLE and console is not None:
        console.print(Panel(body or "(no content)", title=f"{title} • {timestamp}", border_style=color))
    else:  # Graceful fallback if Rich isn't installed
        print(f"[{timestamp}] {title}:\n{body}\n{'-'*40}")

def _print_tool_event(kind: str, node_name: str, name: str, args=None, result: str | None = None) -> None:
    color = _AGENT_COLOR_MAP.get("tool_call" if kind == "call" else "tool_result", "bright_blue")
    timestamp = datetime.now().strftime("%H:%M:%S")
    title = f"TOOL {kind.upper()} • {name} • {node_name}"
    desc = ""
    if kind == "call":
        desc = f"args: {args!r}" if args is not None else "(no args)"
    else:
        desc = result or "(no result)"
    if _RICH_AVAILABLE and console is not None:
        console.print(Panel(desc, title=f"{title} • {timestamp}", border_style=color))
    else:
        print(f"[{timestamp}] {title}:\n{desc}\n{'-'*40}")

############################################################### MULTI AGENT COLLABORATION CONSOLE MONITORING ###############################################################

############################################################### AGENT DEFINITION AND CONFIGURATION SECTION ###############################################################
'''
This section is to define and configure the agents.
1. AI Tutor Agent
- Mainly communicate with user and assign task to other agent.

2. Planner Agent
- Mainly generate study plan that suits user needs.

3. Learning Materials Agent
- Mainly generate resources that user can use to learn.
'''
def _build_planner_agent():
    prompt = (
        """
        <role>
        You are a study plan agent.
        </role>

        <instructions>
        1. Think carefully how to draft a learning plan based on syllabus provided.
        2. List out the thinking steps on how would you construct the learning plan, don't include this steps in study plan.
        3. The syllabus consist of several chapters and each chapter will have sub-topics.
        4. Each chapter needs to list out estimated time and study objectives.
        5. The objective needs to provide a 5 sentences overview that describe the entire topic which include the sub-topic also.
        6. Follow the output format provided strictly while drafting study plan.
        7. Once the study plan is generated, pass it to AI Tutor Agent for review.
        </instructions>

        <Output format>
        <subject name>
        Cloud Computing and Big Data Study Plan
        </subject name>
        <Topic Overview>
        1.	Introduction to Cloud Computing  
        2.	Classification and Business Case  
        3.	Designing Cloud Applications    
        </Topic Overview>
        <topic 1>
        <Topic Name>
        Introduction to Cloud Computing
        </Topic Name>
        <Topic Overview>
        This topic provides foundational knowledge of cloud computing, including its core principles, how virtualization supports cloud infrastructure, and how virtual machines (VMs) function as part of scalable and flexible cloud services.
        </Topic Overview>
        <Sub-Topic Name>
        Primary definition and principles of cloud computing
        </Sub-Topic Name>
        <Study-Objectives>
        Understand the core concept of cloud computing, including key characteristics such as on-demand self-service, broad network access, resource pooling, rapid elasticity, and measured service.
        </Study-Objectives>
        <Estimated time>
        1 Hour
        </Estimated time>
        <Sub-Topic Name>
        Basic technological principles of cloud computing
        </Sub-Topic Name>
        <Study-Objectives>
        Explore the underlying technologies that enable cloud computing, including distributed systems, networking, storage systems, and resource abstraction.
        </Study-Objectives>
        <Estimated time>
        1 Hour
        </Estimated time>
        <Sub-Topic Name>
        Virtualization essentials and working with VMs
        </Sub-Topic Name>
        <Study-Objectives>
        Learn about virtualization technologies, hypervisors, and how virtual machines are created and managed to provide scalable and isolated computing environments in cloud platforms.
        </Study-Objectives>
        <Estimated time>
        2 Hours
        </Estimated time>
        </topic 1>
        </output format>
        """
    )
    return create_react_agent(
        model=llm,
        tools = [],
        prompt=prompt,
        name="planner_agent",
    )

def _build_learning_materials_agent(preference_label: str):
    """Build the learning materials agent with the user's learning preference injected.

    preference_code: one of ('beginner', 'intermediate', 'advanced')
    preference_label: human-friendly label ('Beginner', 'Intermediate', 'Advanced')
    """
    # Guidance to tailor outputs to the student's level
    level_guidance = {
        "beginner": (
            "- Prefer foundational explanations, avoid heavy jargon, include analogies and step-by-step guidance.\n"
            "- Share entry-level resources (intro videos, basic tutorials, gentle documentation).\n"
        ),
        "intermediate": (
            "- Assume basic familiarity; balance concepts with hands-on examples.\n"
            "- Provide resources that deepen understanding (labs, projects, mid-level articles).\n"
        ),
        "advanced": (
            "- Be concise and technical; emphasize depth, trade-offs, and best practices.\n"
            "- Provide advanced resources (whitepapers, specs, research posts, challenging projects).\n"
        ),
    }

    guidance = level_guidance.get(level_guidance[preference_label.lower()])  # safe default

    prompt = (
        """
        <role>
        You are a personalized learning materials agent. Your role is to generate high-quality, tailored learning content for students based on their skill level and preferences (e.g., beginner, intermediate, advanced).
        </role>
        """ 
        "<instructions> \n"
        "Proceed the instruction step by step accordingly. \n"
        "1. Analyze the study plan received from supervisor."
        f"2. According to the user's learning preferences {preference_label}, including their study objectives and any specified format preferences (text, videos, etc.). \n"
        f"3. Follow the guidance : {guidance} to prepare learning materials. \n"
        "4. List out a simple plan on how would you prepare the learning materials and construct the query to search information you needed. \n"
        "5. Follow the plan you just drafted, use the web search tool provided to gather reliable and up-to-date resources to support your content development. Prioritize trustworthy sources such as academic articles, educational websites, and YouTube educational videos. \n"
        "6. Include all referenced articles and media resources inside a `<references>` section. \n" 
        "7. The structure of knowledge will include topic, sub-topic, overview, explanation. \n"
        "8. Provide an overview what will be included in later explanation. \n"
        "9. Explanation have to be comprehensive, details, and in-depth. \n"
        "10. For each sub-topic, explanation length have to be between 10 to 15  paragraphs. Each paragraph need to have at least 10 sentences. \n"
        "11. Your explanation output format need to follow the sample output mentioned in <sample output> tag. \n"
        "12. After you generate the explanations, use draft_notes tool to write it in docx. \n"
        "13. Return the download links to supervisor. \n"
        "</intsructions> \n"
        """
        <sample output>
        <topic> Topic: Introduction to Cloud Computing </topic>
        <overview>
        Cloud computing is a modern approach to delivering computing services—such as storage, servers, applications, and networking—over the internet, allowing users to access resources on demand without owning physical infrastructure. 
        It offers flexibility, scalability, and cost-efficiency, enabling businesses and individuals to deploy and manage IT services more efficiently. 
        By leveraging key technologies like virtualization and models such as IaaS, PaaS, and SaaS, cloud computing supports a wide range of use cases from web hosting to data analytics. As organizations increasingly adopt cloud solutions, understanding its principles and technologies is essential for navigating today's digital landscape.
        </overview>
        <explanation>
        <sub-topic> Primary definition and principles of cloud computing </sub-topic> 
        <paragraph>
        Cloud computing refers to the delivery of computing services—including servers, storage, databases, networking, software, analytics, and intelligence—over the internet (“the cloud”) to offer faster innovation, flexible resources, and economies of scale. 
        Instead of owning and maintaining physical data centers or servers, users can access technology services on demand from cloud providers like Amazon Web Services (AWS), Microsoft Azure, and Google Cloud Platform. 
        This on-demand model reduces the complexity and cost of owning infrastructure and managing software, making it accessible for both businesses and individual users.
        </paragraph>
        <paragraph>
        One of the core principles of cloud computing is resource pooling, which allows cloud providers to serve multiple customers using a multi-tenant model, dynamically assigning and reallocating resources based on demand. 
        Another key principle is broad network access, meaning services are available over the network and accessed through standard mechanisms that promote use across various devices such as phones, tablets, laptops, and desktops. 
        Cloud computing also embraces measured services, allowing for transparent monitoring, control, and optimization of resources, which ensures efficient service provisioning and usage-based billing.
        </paragraph>
        <paragraph>
        Cloud computing models are typically divided into three service types: Infrastructure as a Service (IaaS), Platform as a Service (PaaS), and Software as a Service (SaaS). 
        Each model offers a different level of control, flexibility, and management, depending on user needs. 
        Additionally, deployment models—public cloud, private cloud, and hybrid cloud—determine how cloud services are hosted and managed. 
        Together, these principles and models define the foundation of modern cloud computing, enabling scalability, reliability, and agility in IT operations.
        </paragraph>
        <sub-topic> Basic technological principles of cloud computing: Virtualization essentials and working with VMs </sub-topic>
        <paragraph>
        Cloud computing relies heavily on the principle of virtualization, which allows physical computing resources such as servers, storage, and networking to be divided into multiple virtual units. 
        Virtualization is achieved through hypervisors—software that enables multiple virtual machines (VMs) to run on a single physical machine while sharing its resources. 
        This technology provides the foundation for cloud computing by enabling better utilization of hardware, isolation of applications, and scalability in managing workloads.
        </paragraph>
        <paragraph>
        Virtual machines are central to this approach. 
        Each VM operates as an independent environment with its own operating system and applications, even though they share the same physical hardware. 
        VMs can be easily created, modified, or deleted, making them ideal for dynamic and on-demand computing environments like the cloud. 
        This flexibility allows cloud providers to offer Infrastructure as a Service (IaaS), where users can deploy and manage VMs based on their specific needs, paying only for what they use.
        </paragraph>
        </explanation>
        <references>
        https://www.ibm.com/think/topics/cloud-computing
        https://www.techtarget.com/searchcloudcomputing/definition/cloud-computing
        https://youtu.be/mxT233EdY5c?si=c9KaWbSkmhoNlWK4
        </references>
        </sample output>
        """
    )
    return create_react_agent(
        model=llm,
        tools = [_web_search, _draft_notes],
        prompt=prompt,
        name="learning_materials_agent",
    )

def _build_supervisor(worker_agent_1, worker_agent_2, memory: AsyncPostgresSaver):
    supervisor_prompt = (
        """
        <role>
        You are the AI Tutor, the primary agent responsible for interacting with the student. You serve as both an academic assistant and a coordinator for two support agents:

        - **Planner Agent**: Generates structured and personalized study plans based on a syllabus.
        - **Learning Materials Agent**: Creates topic-specific learning notes based on the finalized study plan.

        You must not perform planning or materials generation tasks yourself — always delegate them to the appropriate support agent.
        </role>

        <instructions>
        Your core responsibilities include:

        1. **Introducing Yourself**
        - Greet the student and briefly explain your role and how you can assist.

        2. **Handling User Queries**
        - Based on the user's input, choose one of the following actions:

        <action 1: Academic Support>
        - If the user asks an academic question or needs help understanding a concept, provide a clear and concise explanation tailored to their level.
        - Use examples, visuals (if supported), and step-by-step reasoning when appropriate.

        <action 2: Study Plan Workflow>
        If the user provides a syllabus and wants a study plan and learning materials, follow this scenario exactly:

        **Scenario: Syllabus Provided**
        1. Forward the received syllabus to the **Planner Agent** to generate a study plan.
        2. Once the Planner Agent reports completion, request the generated study plan.
        3. Present the study plan to the user for review.
        4. Wait for user confirmation:
        - If **approved**, use the `write_study_plan` tool to generate a `.docx` version of the plan.
        - If **not approved**, request a refinement from the Planner Agent and repeat the review process.
        5. After final approval, extract the first topic (or the next topic as per user preference) from the study plan.
        6. Send this topic to the **Learning Materials Agent** to generate learning notes.
        7. When the Learning Materials Agent reports completion, retrieve the learning materials and provide the **download URL** to the user.
        </action 2>
        </instructions>

        <rules>
        Strict Rules:
        - Do **not** perform planning or content creation tasks yourself.
        - Always delegate one agent task at a time.
        - Do **not** modify the format of any output received from support agents.
        </rules>
        """
    )

    compiled = create_supervisor(
        model=llm,
        agents=[worker_agent_1, worker_agent_2],
        tools=[_write_study_plan],
        prompt=supervisor_prompt,
        add_handoff_back_messages=True,
        output_mode="full_history",
    ).compile(checkpointer=memory)
    return compiled

############################################################### AGENT DEFINITION AND CONFIGURATION SECTION ###############################################################


############################################################### AGENT WORKFLOW IMPLEMENTATION ###############################################################
async def run_supervised_agents(user_message: str, thread_id: Optional[str], pool: AsyncConnectionPool) -> str:
    """
    Run the multi-agent system with a supervisor and return only the final user-facing response.
    - Uses a Postgres-backed memory tied to the provided thread_id.
    - Hides internal agent chatter; only the supervisor's final AIMessage is returned.
    """
    # Fallback thread id if not provided
    effective_thread_id = thread_id

    planner_agent = _build_planner_agent()

    # Fetch user's learning preference from the conversation thread (safe default to Beginner)
    preference_label = "beginner"
    try:
        if thread_id:
            # Deferred import to avoid circulars at module import time
            from users.models import User
            # ORM call in async context via sync_to_async
            preference_label = await sync_to_async(User.objects.get(thread_id=thread_id)).learning_preference
            
            logger.info(f"Learning preference resolved for thread {thread_id}: {preference_label})")
    except Exception as e:
        logger.warning(f"Failed to resolve learning preference for thread {thread_id}: {e}")

    learning_materials_agent = _build_learning_materials_agent(preference_label)

    # Prepare containers for results and transcript
    final_response: Optional[str] = None

    # Use the provided async connection pool from the caller (consumers.py)
    async with pool.connection() as conn:
        memory = AsyncPostgresSaver(conn)

        # Instantiate supervisor with memory
        supervisor = _build_supervisor(planner_agent, learning_materials_agent, memory)

        logger.info("Supervisor started with provided DB connection pool")
        # Try async stream first (if supported); else fall back to sync stream
        async for chunk in supervisor.astream(
            {"messages": [HumanMessage(content=user_message)]},
            {"configurable": {"thread_id": effective_thread_id}},
        ):
            # Print all node emissions for full collaboration visibility
            for node_name, payload in chunk.items():

                messages = payload.get("messages", []) if isinstance(payload, dict) else []

                i = 0
                for m in messages:
                    try:
                        _print_collab_message(node_name, m)
                    except Exception:
                        # Don't let logging break the flow
                        pass
                    # Print tool call/result events if present
                    try:
                        if hasattr(m, "tool_calls") and m.tool_calls:
                            for tc in m.tool_calls:
                                name = tc.get("name") if isinstance(tc, dict) else getattr(tc, "name", None)
                                args = tc.get("args") if isinstance(tc, dict) else getattr(tc, "args", None)
                                _print_tool_event("call", m.name, name or "(unknown)", args=args) 
                        if (ToolMessage is not None) and isinstance(m, ToolMessage):
                            _print_tool_event("result", messages[i-1].name, getattr(m, "name", "tool"), result=_extract_text_content(getattr(m, "content", "")))
                    except Exception:
                        pass
                    if node_name == "supervisor" and isinstance(m, AIMessage):
                        # Track latest supervisor reply to return to user
                        final_response = m.content

                    i+=1
            # final_response = chunk["supervisor"]["messages"][-1].content
        # Sync fallback omitted for brevity; async is preferred

    # Ensure we return something even if no supervisor AIMessage was produced
    if not final_response:
        final_response = "I'm here to help. Could you please rephrase or provide more details?"

    return final_response
############################################################### AGENT WORKFLOW IMPLEMENTATION ###############################################################